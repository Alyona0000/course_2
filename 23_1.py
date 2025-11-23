import os
import docx

def get_para_data(output_doc_name, paragraph, paragraph2):
    output_para = output_doc_name.add_paragraph()

    for run, run2 in zip(paragraph.runs, paragraph2.runs):
        output_run = output_para.add_run(run.text)
        output_run.bold = run2.bold
        output_run.italic = run2.italic
        output_run.underline = run2.underline
        output_run.font.name = run2.font.name
        output_run.font.size = run2.font.size
        output_run.font.color.rgb = run2.font.color.rgb
        output_run.style.name = run2.style.name

    output_para.paragraph_format.alignment = paragraph2.paragraph_format.alignment

if __name__ == '__main__':
    path = '111' #input('Enter path to directory: ')
    files = os.listdir(path)
    files = {i:fname for i, fname in zip(range(1, len(files) + 1), files)}
    message = 'Choose reference file: {}\n'.format(''.join([f'\n\t{i} - {fname}' for i, fname in files.items()]))
    file = files[int(input(message))]
    mainFile = docx.Document(rf'{path}\{file}')
    docFiles = set(os.listdir(path)) - {file}

    for docFile in docFiles:
        doc = docx.Document()

        for par, par2 in zip(docx.Document(rf'{path}\{docFile}').paragraphs, mainFile.paragraphs):
            get_para_data(doc, par, par2)

        doc.save(rf'111\{docFile}') # doc.save(rf'{path}\{docFile}')