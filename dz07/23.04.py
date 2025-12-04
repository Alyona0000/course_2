# T23.4 У документі Word містяться дати у форматі dd.mm.yyyy або
# підкреслення __.__.____.
# Знайти всі дати у тексті.
# Замість підкреслень вставити поточну дату. Зберегти оновлений документ.

import re
import os
from datetime import datetime
from docx import Document

# ----------------------- НАЛАШТУВАННЯ ШЛЯХІВ -----------------------

input_file = r"C:\Users\alenk\Repository\Course_2\input.docx"
output_file = r"C:\Users\alenk\Repository\Course_2\output.docx"
folder = r"C:\Users\alenk\Repository\Course_2"

# ----------------------- ДІАГНОСТИКА -----------------------

print("=== ВМІСТ ПАПКИ ===")
try:
    files = os.listdir(folder)
    for f in files:
        print(" -", f)
except FileNotFoundError:
    print("❌ Папку не знайдено! Перевір шлях:", folder)
    exit()

print("\n=== ПЕРЕВІРКА ІСНУВАННЯ ФАЙЛУ ===")
file_exists = os.path.exists(input_file)
print("Файл існує?", file_exists)

if not file_exists:
    print("❌ Файл не знайдений. Перевір назву і те, чи знаходиться в папці.")
    exit()

print("\n=== СПРОБА ВІДКРИТИ ФАЙЛ ЯК DOCX ===")
try:
    test_doc = Document(input_file)
    print("✔ Файл успішно відкрився як справжній DOCX!")
except Exception as e:
    print("❌ Не вдалося відкрити файл!")
    print("Тип помилки:", type(e).__name__)
    print("Текст:", e)
    print("\nЙмовірні причини:")
    print("- Файл пошкоджений")
    print("- Файл не .docx (а .doc або перейменований .txt)")
    print("- Порожній файл зі зміненим розширенням")
    exit()

# ----------------------- ОСНОВНА ФУНКЦІЯ -----------------------

def process_dates(input_file, output_file):
    doc = Document(input_file)

    date_pattern = re.compile(r"\b\d{2}\.\d{2}\.\d{4}\b")
    blank_pattern = re.compile(r"__\.__\.____")
    today = datetime.today().strftime("%d.%m.%Y")

    for paragraph in doc.paragraphs:
        text = paragraph.text
        new_text = blank_pattern.sub(today, text)
        found_dates = date_pattern.findall(new_text)
        paragraph.text = new_text

    doc.save(output_file)
    print("\n✔ Готово! Оновлений файл збережено як:", output_file)

# ----------------------- ЗАПУСК -----------------------

process_dates(input_file, output_file)
